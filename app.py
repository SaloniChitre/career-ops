import streamlit as st
import json
import pandas as pd
import pdfplumber
from pathlib import Path
from io import BytesIO
from docx import Document
import re
from streamlit_lottie import st_lottie
import requests
from bs4 import BeautifulSoup

def load_lottieurl(url):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except Exception:
        return None

# --- NEW: URL PARSER HELPER ---
def scrape_jd_from_url(url):
    """
    Enhanced JD Scraper with Login Wall Detection and Header Spoofing.
    Designed to prevent 'weird' LinkedIn login text from being treated as a JD.
    """
    try:
        # 1. Spooking headers to look like a real modern browser
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,webp,image/apng,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9",
            "Referer": "https://www.google.com/"
        }
        
        # Use a session to handle potential redirects and basic cookies
        session = requests.Session()
        response = session.get(url, headers=headers, timeout=12)
        
        # 2. Check for immediate LinkedIn Authwall redirect
        if "linkedin.com/checkpoint" in response.url or "authwall" in response.url:
            return "⚠️ LinkedIn Login Wall detected. Automated scraping is blocked by LinkedIn's security. Please copy/paste the JD text manually into the box below."
            
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 3. Clean the HTML (Remove noise like scripts, navbars, and footers)
            for element in soup(["script", "style", "nav", "footer", "header", "noscript"]):
                element.decompose()
            
            # Get text and clean up whitespace
            text = soup.get_text(separator=' ')
            lines = (line.strip() for line in text.splitlines())
            clean_text = "\n".join(line for line in lines if line)
            
            # 4. LOGIN WALL VALIDATION (The "Weird Text" Filter)
            # We look for common login-page keywords that shouldn't be in a real JD
            login_keywords = [
                "Sign in", "Join now", "User Agreement", "Privacy Policy", 
                "Cookie Policy", "Forgot password?", "Email or phone",
                "Keep me logged in", "New to LinkedIn?"
            ]
            
            # Count how many login keywords appeared
            matches = [word for word in login_keywords if word.lower() in clean_text.lower()]
            
            # If we find more than 3 login-related terms, it's almost certainly a login wall
            if len(matches) >= 3 or len(clean_text) < 300:
                return "⚠️ Security Wall Detected: The site returned a login screen instead of the job details. \n\nACTION REQUIRED: Please open the link in your browser and manually paste the description below."
            
            # 5. JAVASCRIPT DETECTION
            if "enable JavaScript" in clean_text or "requires JavaScript" in clean_text:
                return "⚠️ JavaScript Wall: This site (likely Ashby or Lever) requires a full browser to render. Please copy/paste the JD manually."
                
            return clean_text
            
        return f"Error: Status Code {response.status_code}"
        
    except Exception as e:
        return f"Error connecting to URL: {str(e)}"

# Load the animation data once
lottie_ai = load_lottieurl("https://assets10.lottiefiles.com/packages/lf20_m6cuL6.json")

# Attempting to import your custom processor logic
try:
    from core.processor import compare_resume_to_job, client
except ImportError:
    client = None
    def compare_resume_to_job(r, j): return {"score": 0, "matching_skills": [], "missing_skills": []}

# --- 1. UI STYLING & CONFIG ---
st.set_page_config(page_title="Deepam's Career Ops", page_icon="✨", layout="wide")

st.markdown("""
    <style>
    .stApp {
        background-color: #0f172a !important;
        background-image: linear-gradient(180deg, #0f172a 0%, #1e1b4b 100%) !important;
    }
    .main * { color: #f1f5f9 !important; }
    [data-testid="stSidebar"] {
        background-color: #1e1b4b !important;
        border-right: 1px solid #334155;
    }
    [data-testid="stSidebar"] * { color: #ffffff !important; }
    [data-testid="stDataFrame"] {
        background-color: #1e293b !important;
        border: 1px solid #334155;
        border-radius: 10px;
    }
    div[data-testid="metric-container"] {
        background-color: #1e293b !important;
        border: 1px solid #4f46e5 !important;
        border-radius: 12px;
        padding: 15px;
    }
    [data-testid="stMetricValue"] { color: #818cf8 !important; }
    header { visibility: hidden; }
    </style>
    """, unsafe_allow_html=True)

LOG_PATH = Path("data/jobs_log.json")

# --- 2. HELPER FUNCTIONS ---
def extract_pdf(file):
    if not file: return None
    try:
        with pdfplumber.open(file) as pdf:
            return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
    except Exception: return None

def save_as_docx(content, title):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def ai_generate(prompt):
    if not client: return "Error: AI Client not configured."
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        return completion.choices[0].message.content
    except Exception as e: return f"Error: {str(e)}"

def update_job_status(job_id, new_status, raw_data):
    job_list = raw_data['jobs'] if isinstance(raw_data, dict) and 'jobs' in raw_data else raw_data
    for job in job_list:
        if str(job.get('id')) == str(job_id):
            job['status'] = new_status.upper()
            break
    with open(LOG_PATH, "w") as f:
        json.dump(raw_data, f, indent=4)
    st.rerun()

# --- 3. UI RENDERING LOGIC ---
def render_job_section(job_list, raw_data, key_suffix):
    if not job_list:
        st.info("No active opportunities found in this category.")
        return

    display_list = []
    is_alert_tab = "alrt" in key_suffix.lower()

    for i, j in enumerate(job_list):
        temp_job = j.copy()
        job_id = temp_job.get('id')
        uid = job_id if job_id and str(job_id) != "None" else f"{key_suffix}_{i}"
        status = str(temp_job.get('status', 'ALERT')).upper()
        
        if is_alert_tab:
            item = {
                "Status": status,
                "Location": temp_job.get('company', 'N/A'), 
                "Company": temp_job.get('title', 'N/A'),    
                "Role": "See JD/Header", 
                "safe_id": uid,
                "header": f"{temp_job.get('title')}", 
                "description": temp_job.get('description', '')
            }
            cols_to_show = ["Status", "Location", "Company", "Role"]
        else:
            item = {
                "Status": status,
                "Role": temp_job.get('title', 'Unknown'),
                "Company": temp_job.get('company', 'N/A'),
                "safe_id": uid,
                "header": f"{temp_job.get('title')} @ {temp_job.get('company')}",
                "description": temp_job.get('description', '')
            }
            cols_to_show = ["Status", "Role", "Company"]
        
        display_list.append(item)

    st.dataframe(pd.DataFrame(display_list)[cols_to_show], use_container_width=True, hide_index=True)

    st.markdown("### 🛠️ Intelligence & Actions")
    
    for job in display_list:
        jid = job['safe_id']
        
        with st.expander(f"💼 {job['header']}"):
            c1, c2 = st.columns([2, 1])
            with c1:
                action = st.selectbox("Select AI Action:", 
                    ["Analyze Fit", "Tailor Resume", "Draft Cover Letter", "Interview Prep"],
                    key=f"act_{jid}")
            with c2:
                current_status = job['Status']
                status_options = ["ALERT", "APPLIED", "IN_PROGRESS", "REJECTED"]
                try:
                    status_index = status_options.index(current_status)
                except ValueError: 
                    status_index = 0
                
                new_stat = st.selectbox("Update Status:", status_options, index=status_index, key=f"stat_{jid}")
                if new_stat != current_status:
                    if st.button("Save Status", key=f"save_{jid}"):
                        update_job_status(jid, new_stat, raw_data)

            st.divider()

            # --- NEW: JOB-SPECIFIC URL PARSER ---
            st.write("**🔗 Auto-Load JD for this Job**")
            job_specific_url = st.text_input("Paste URL:", key=f"url_input_{jid}")
            if st.button("Fetch JD Content", key=f"btn_fetch_{jid}"):
                if job_specific_url:
                    with st.spinner("Parsing..."):
                        # Save specifically for this job ID
                        st.session_state[f'parsed_jd_{jid}'] = scrape_jd_from_url(job_specific_url)
                else:
                    st.warning("Please paste a URL first.")
            
            col_in, col_out = st.columns([1, 1.5])
            
            with col_in:
                st.write("**Job Description Context**")
                
                # Check for job-specific parsed text, then fall back to log description
                current_jd_text = st.session_state.get(f'parsed_jd_{jid}', job['description'])
                
                final_jd = st.text_area(
                    "JD Content:", 
                    value=current_jd_text, 
                    height=250, 
                    key=f"txt_{jid}_{hash(current_jd_text)}" 
                )

            with col_out:
                st.write(f"**AI Results: {action}**")
                current_resume = st.session_state.get('resume_text')

                if not current_resume:
                    st.error("Please upload your resume in the sidebar!")
                elif not final_jd or len(final_jd.strip()) < 20:
                    st.warning("JD Content empty. Paste text or use the Fetcher above.")
                else:
                    if action == "Analyze Fit":
                        if st.button("Run Match Score", key=f"go_fit_{jid}"):
                            with st.status("🔮 Analyzing compatibility...", expanded=True):
                                res = compare_resume_to_job(current_resume, final_jd)
                            score = res.get('score', 0)
                            st.metric("Compatibility", f"{score}%")
                            st.write("**Matching Skills:**", ", ".join(res.get('matching_skills', [])))
                            st.write("**Gaps:**", ", ".join(res.get('missing_skills', [])))
                            if score > 0: st.balloons()

                    elif action == "Tailor Resume":
                        if st.button("Generate Bullets", key=f"go_tailor_{jid}"):
                            with st.spinner("Tailoring..."):
                                prompt = f"Based on my resume: {current_resume[:1500]} and this JD: {final_jd[:1000]}, write 3 high-impact resume bullets."
                                st.write(ai_generate(prompt))

                    elif action == "Draft Cover Letter":
                        if st.button("Generate Letter", key=f"go_cl_{jid}"):
                            with st.spinner("Drafting..."):
                                prompt = f"Write a professional cover letter for {job['header']} using my background: {current_resume[:1500]}."
                                content = ai_generate(prompt)
                                st.write(content)
                                st.download_button("Download Letter", data=save_as_docx(content, "Cover Letter"), file_name=f"CL.docx", key=f"dl_cl_{jid}")

                    elif action == "Interview Prep":
                        if st.button("Generate Prep Guide", key=f"go_prep_{jid}"):
                            with st.spinner("Preparing..."):
                                prompt = f"Generate 5 technical interview questions for {job['header']} based on this JD: {final_jd[:1500]}."
                                st.write(ai_generate(prompt))

# --- 4. MAIN APP ---
st.sidebar.markdown("""
    <div style="text-align: center;">
        <img src="https://api.dicebear.com/7.x/bottts/svg?seed=Deepam" width="100" style="border-radius: 50%; background: #ddd6fe; padding: 10px; border: 2px solid #8b5cf6;">
        <h3 style="color: #6d28d9; margin-bottom: 20px;">Deepam's AI Co-Pilot</h3>
    </div>
    """, unsafe_allow_html=True)

st.sidebar.divider()
st.sidebar.title("👤 Candidate Profile")
uploaded_resume = st.sidebar.file_uploader("Upload Resume (PDF)", type=["pdf"])

if uploaded_resume:
    st.session_state['resume_text'] = extract_pdf(uploaded_resume)

st.title("🚀 Career Ops Command Center")

if LOG_PATH.exists():
    with open(LOG_PATH, "r") as f:
        raw_data = json.load(f)
    data_list = raw_data['jobs'] if isinstance(raw_data, dict) and 'jobs' in raw_data else raw_data

    def get_by_status(status_str):
        return [j for j in data_list if isinstance(j, dict) and str(j.get('status', '')).upper() == status_str.upper()]

    applied_list = get_by_status('APPLIED')
    prog_list = get_by_status('IN_PROGRESS')
    alert_list = get_by_status('ALERT')
    rej_list = get_by_status('REJECTED')

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total Apps", len(applied_list))
    m2.metric("Interviews", len(prog_list))
    m3.metric("New Alerts", len(alert_list))
    success_rate = round((len(prog_list) / max(len(applied_list), 1)) * 100)
    m4.metric("Success Rate", f"{success_rate}%")

    tabs = st.tabs(["📤 Applied", "👀 In Progress", "🔔 Alerts", "❌ Rejected"])
    with tabs[0]: render_job_section(applied_list, raw_data, "app")
    with tabs[1]: render_job_section(prog_list, raw_data, "prog")
    with tabs[2]: render_job_section(alert_list, raw_data, "alrt")
    with tabs[3]: 
        if rej_list:
            rej_df = pd.DataFrame(rej_list)
            st.dataframe(rej_df[["title", "company"]], use_container_width=True, hide_index=True)
        else:
            st.info("No rejections yet.")
else:
    st.warning("No data found. Please run the sync script.")